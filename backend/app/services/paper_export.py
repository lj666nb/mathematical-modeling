"""
============================================================
论文导出服务 — Markdown → DOCX 转换（Pandoc + LaTeX公式渲染）
============================================================

使用 Pandoc 将 Markdown 转为 Word .docx：
  - $...$ / $$...$$ → Word 原生方程对象 (OMML)，公式正确渲染
  - Markdown 表格 → Word 表格
  - 代码块 → 格式化代码
  - 标题层级 → Word Heading 样式
"""
from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _style_code_blocks(doc):
    """
    🆕 美化 DOCX 中的代码块：
    - 查找 Pandoc 生成的 Source Code 段落
    - 添加左边框线（蓝色/灰色竖线）
    - 设置 Consolas 等宽字体
    - 浅灰色背景
    - 适当缩进
    """
    CODE_BG = "F5F5F5"       # 浅灰背景
    BORDER_COLOR = "4472C4"  # 蓝色边框
    FONT_NAME = "Consolas"
    FONT_SIZE = Pt(9)        # 代码字体大小

    # Try to find/create Source Code style
    try:
        code_style = doc.styles['Source Code']
    except KeyError:
        try:
            code_style = doc.styles.add_style('Source Code', 1)  # WD_STYLE_TYPE.PARAGRAPH = 1
        except Exception:
            # Fallback: style paragraphs individually
            code_style = None

    if code_style:
        code_style.font.name = FONT_NAME
        code_style.font.size = FONT_SIZE
        code_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        code_style.paragraph_format.space_before = Pt(2)
        code_style.paragraph_format.space_after = Pt(2)
        code_style.paragraph_format.line_spacing = 1.2
        code_style.paragraph_format.left_indent = Cm(0.5)

    # 🆕 Process each paragraph to add borders and shading
    in_code_block = False
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""

        # Detect code paragraphs
        is_code = (
            style_name == 'Source Code' or
            'Source Code' in style_name or
            'code' in style_name.lower()
        )

        # Also detect by font — Pandoc may set Consolas/Courier on inline code
        if not is_code and para.runs:
            for run in para.runs:
                if run.font.name and any(
                    m in run.font.name.lower()
                    for m in ('consolas', 'courier', 'monaco', 'monospace', 'source code')
                ):
                    is_code = True
                    break

        if is_code:
            in_code_block = True
            _add_paragraph_border(para, BORDER_COLOR, CODE_BG)
            # Ensure all runs use monospace
            for run in para.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
                if not run.font.color.rgb:
                    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        elif in_code_block and para.text.strip() == '':
            # Empty line after code block — end of block
            in_code_block = False


def _add_paragraph_border(para, border_color: str, bg_color: str):
    """Add left border + background shading to a single paragraph (code block line)"""
    pPr = para._element.get_or_add_pPr()

    # Background shading
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), bg_color)
    shd.set(qn('w:val'), 'clear')
    pPr.insert(0, shd)

    # Paragraph borders
    pBdr = OxmlElement('w:pBdr')

    # Left border (thick colored line)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')  # border width
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), border_color)
    pBdr.append(left)

    # Top border (thin)
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'D0D0D0')
    pBdr.append(top)

    # Bottom border (thin)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D0D0D0')
    pBdr.append(bottom)

    # Right border (thin)
    right = OxmlElement('w:right')
    right.set(qn('w:val'), 'single')
    right.set(qn('w:sz'), '4')
    right.set(qn('w:space'), '4')
    right.set(qn('w:color'), 'D0D0D0')
    pBdr.append(right)

    pPr.append(pBdr)


def _preprocess_markdown(md_text: str, figures_dir: Path) -> str:
    """
    预处理 Markdown，确保 Pandoc 能正确解析：
    1. ####+ → ###（Word 最多3级标题）
    2. 图片路径 → 绝对路径（Pandoc 需要能找到）
       🆕 增强：精确匹配失败时，尝试模糊匹配（按 question_id/关键词）
    3. 清理反引号包裹的公式（转为 $...$ 包裹）
    """
    # 1. 标题降级
    md_text = re.sub(r'^#{4,6}\s+', '### ', md_text, flags=re.MULTILINE)

    # 🆕 预扫描：建立 figures_dir 下所有图片的索引
    available_images: dict[str, Path] = {}
    if figures_dir.exists():
        for img_file in figures_dir.glob("*.png"):
            available_images[img_file.name.lower()] = img_file
        for img_file in figures_dir.glob("*.jpg"):
            available_images[img_file.name.lower()] = img_file

    def _find_best_image(rel_path: str) -> Path | None:
        """多策略查找图片文件（LLM 经常编造文件名，需要健壮的模糊匹配）"""
        filename = Path(rel_path).name
        fname_lower = filename.lower()

        # 策略1: 精确匹配文件名
        if fname_lower in available_images:
            return available_images[fname_lower]

        # 策略2: 提取问题编号 Q1/Q2/Q3，找该问题的所有图
        q_match = re.search(r'[qQ](\d+)', fname_lower)
        if q_match:
            q_num = q_match.group(1)
            q_candidates = sorted(
                [p for name, p in available_images.items()
                 if f'q{q_num}' in name.lower() or f'fig_q{q_num}' in name.lower()],
                key=lambda p: p.stat().st_size, reverse=True  # 最大的图最可能是路径/结果图
            )
            if q_candidates:
                return q_candidates[0]

        # 策略3: 关键词 → 可用文件名匹配
        keyword_to_file_patterns = {
            "cost": ["cost", "pie", "bar", "fig_q"],
            "route": ["route", "spatial", "map", "path", "fig_q"],
            "converge": ["converge", "curve", "line", "fig_q"],
            "sensitivity": ["sensitivity", "radar", "fig_q"],
            "radar": ["radar", "sensitivity", "fig_q"],
            "overview": ["overview"],
            "correlation": ["correlation", "heatmap"],
            "heatmap": ["heatmap", "correlation"],
            "distrib": ["distribution", "dist", "hist"],
            "missing": ["missing"],
            "comparison": ["comparison", "contrast", "fig_q"],
            "pie": ["pie", "fig_q"],
            "map": ["map", "spatial", "fig_q"],
            "path": ["path", "spatial", "route", "fig_q"],
            "curve": ["curve", "converge", "line", "fig_q"],
        }
        for kw, patterns in keyword_to_file_patterns.items():
            if kw in fname_lower:
                for pat in patterns:
                    candidates = sorted(
                        [p for name, p in available_images.items() if pat in name.lower()],
                        key=lambda p: p.stat().st_size, reverse=True
                    )
                    if candidates:
                        return candidates[0]

        # 策略4: 最后的回退 — 返回第一张 overview 图或任何可用的图
        for name, p in available_images.items():
            if 'overview' in name.lower():
                return p
        # 返回任意可用图
        if available_images:
            return sorted(available_images.values(), key=lambda p: p.stat().st_size, reverse=True)[0]

        return None

    # 2. 图片路径 → 绝对路径（多策略查找）
    def _fix_img_path(m):
        alt = m.group(1)
        rel_path = m.group(2)
        # 策略1: 直接从 figures_dir 精确匹配文件名
        abs_path = (figures_dir / Path(rel_path).name).resolve()
        if abs_path.exists() and abs_path.stat().st_size > 100:
            return f'![{alt}]({abs_path.as_posix()})'
        # 策略2: 模糊匹配
        best = _find_best_image(rel_path)
        if best is not None:
            return f'![{alt}]({best.resolve().as_posix()})'
        # 策略3: 如果没有任何匹配，返回原样（不删除引用）
        return m.group(0)

    md_text = re.sub(r'!\[(.*?)\]\((.*?)\)', _fix_img_path, md_text)

    # 3. 反引号公式 → $...$（检测含 LaTeX 模式的反引号内容）
    def _convert_backtick_math(m):
        code = m.group(1)
        if re.search(r'[_^\\]|\\sum|\\frac|\\int|\\prod|\\cdot|\\times|\\leq|\\geq|\\in\\b|\\forall|\\exists|\\alpha|\\beta|\\gamma|\\delta|\\epsilon|\\theta|\\lambda|\\mu|\\pi|\\sigma|\\omega|\\Sigma|\\Omega', code):
            return f'${code}$'
        return m.group(0)

    md_text = re.sub(r'`([^`]+?)`', _convert_backtick_math, md_text)

    return md_text


def _add_page_number(doc):
    """添加页码"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run._r.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)


def _postprocess_docx(docx_path: Path) -> None:
    """
    Post-process the Pandoc-generated DOCX:
    - Set page margins (A4 standard)
    - Set default Chinese fonts (宋体 for body, 黑体 for headings)
    - Add page numbers
    - 🆕 Style code blocks: borders, monospace font, gray background, line numbers
    """
    doc = Document(str(docx_path))

    # Page setup
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # Normal style
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # Heading styles
    for level in range(1, 4):
        try:
            h_style = doc.styles[f'Heading {level}']
            h_style.font.name = '黑体'
            h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            h_style.font.color.rgb = RGBColor(0, 0, 0)
            if level == 1:
                h_style.font.size = Pt(16)
            elif level == 2:
                h_style.font.size = Pt(14)
            else:
                h_style.font.size = Pt(13)
        except KeyError:
            pass

    # 🆕 Style code blocks — Pandoc uses "Source Code" style for fenced code blocks
    _style_code_blocks(doc)

    # Page numbers
    _add_page_number(doc)

    doc.save(str(docx_path))


def convert_md_to_docx(
    md_text: str,
    figures_dir: Path,
    output_path: Path,
    title: str = "数学建模论文"
) -> str:
    """
    将 Markdown 论文转换为 Word 文档 (.docx)

    使用 Pandoc 进行转换，LaTeX 公式 ($...$ / $$...$$) 自动渲染为
    Word 原生方程对象 (OMML)。

    Args:
        md_text: Markdown 论文内容
        figures_dir: 图表文件所在目录
        output_path: 输出 .docx 路径
        title: 论文标题

    Returns:
        输出文件路径字符串
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Pre-process markdown
    md_text = _preprocess_markdown(md_text, figures_dir)

    # Write pre-processed markdown to temp file
    with tempfile.NamedTemporaryFile(
        mode='w', suffix='.md', encoding='utf-8', delete=False
    ) as tmp_md:
        tmp_md.write(md_text)
        tmp_md_path = Path(tmp_md.name)

    try:
        # Run pandoc: MD → DOCX
        # --from markdown+tex_math_dollars: support $...$ and $$...$$ math
        # --to docx: Word output
        pandoc_cmd = [
            'pandoc',
            str(tmp_md_path),
            '-o', str(output_path),
            '--from', 'markdown+tex_math_dollars+pipe_tables+fenced_code_blocks',
            '--to', 'docx',
            '--wrap=none',
        ]

        result = subprocess.run(
            pandoc_cmd,
            capture_output=True,
            text=True,
            timeout=120,
        )

        if result.returncode != 0:
            # Pandoc failed — fall back to basic python-docx conversion
            _fallback_convert(md_text, figures_dir, output_path, title)
        else:
            # Post-process the pandoc-generated DOCX
            _postprocess_docx(output_path)

    except (FileNotFoundError, subprocess.TimeoutExpired):
        # Pandoc not available or timed out — fallback
        _fallback_convert(md_text, figures_dir, output_path, title)
    finally:
        # Clean up temp file
        try:
            tmp_md_path.unlink()
        except Exception:
            pass

    return str(output_path)


def _fallback_convert(
    md_text: str,
    figures_dir: Path,
    output_path: Path,
    title: str = "数学建模论文"
) -> str:
    """
    Fallback: python-docx conversion when pandoc is unavailable.
    🆕 增强：支持图片嵌入、代码块格式化、表格渲染
    """
    from docx import Document as Doc
    from docx.shared import Inches, Pt as Pt2, Cm as Cm2, RGBColor as RGB2
    from docx.enum.text import WD_ALIGN_PARAGRAPH as WD2

    doc = Doc()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm2(21)
    section.page_height = Cm2(29.7)
    section.top_margin = Cm2(2.54)
    section.bottom_margin = Cm2(2.54)
    section.left_margin = Cm2(3.17)
    section.right_margin = Cm2(3.17)

    # Styles
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt2(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt2(6)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for lv in range(1, 4):
        h_style = doc.styles[f'Heading {lv}']
        h_style.font.name = '黑体'
        h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        h_style.font.color.rgb = RGB2(0, 0, 0)
        if lv == 1:
            h_style.font.size = Pt2(16)
        elif lv == 2:
            h_style.font.size = Pt2(14)
        else:
            h_style.font.size = Pt2(13)

    _add_page_number(doc)

    # 🆕 Pre-build image index for fuzzy matching
    available_images: dict[str, Path] = {}
    if figures_dir.exists():
        for img_file in figures_dir.glob("*.png"):
            available_images[img_file.name.lower()] = img_file

    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_lines: list[str] = []

    def _flush_code_block():
        """Write accumulated code lines as formatted paragraphs"""
        nonlocal code_lines, in_code_block
        if not code_lines:
            return
        # Find code language from opening fence
        lang = ""
        if code_lines[0].startswith("```"):
            lang = code_lines[0][3:].strip()
            code_lines = code_lines[1:]
        if code_lines and code_lines[-1].startswith("```"):
            code_lines = code_lines[-1][3:].strip() and code_lines[:-1] or code_lines[:-1]

        for cl in code_lines:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt2(1)
            p.paragraph_format.space_after = Pt2(1)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(cl if cl else " ")
            run.font.name = 'Consolas'
            run.font.size = Pt2(9)
            run.font.color.rgb = RGB2(0x1A, 0x1A, 0x2E)
        code_lines = []
        in_code_block = False

    def _find_image_path(rel_path: str) -> Path | None:
        """Find image file by name with fuzzy fallback"""
        fname = Path(rel_path).name.lower()
        if fname in available_images:
            return available_images[fname]
        # Fuzzy: match by question number
        qm = re.search(r'[qQ](\d+)', fname)
        if qm:
            qp = f"fig_q{qm.group(1)}"
            candidates = sorted([p for n, p in available_images.items() if qp in n], key=lambda p: p.name)
            if candidates:
                return candidates[0]
        return None

    while i < len(lines):
        line = lines[i]

        # 🆕 Code block detection
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code_block:
                # End of code block
                code_lines.append(line)
                _flush_code_block()
                i += 1
                continue
            else:
                # Start of code block
                in_code_block = True
                code_lines.append(line)
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Skip empty
        if not stripped:
            i += 1
            continue

        # 🆕 Image detection
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', stripped)
        if img_match:
            alt = img_match.group(1)
            rel_path = img_match.group(2)
            img_path = _find_image_path(rel_path)
            if img_path is None:
                # Try figures_dir / filename
                img_path = figures_dir / Path(rel_path).name

            if img_path and img_path.exists() and img_path.stat().st_size > 100:
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(img_path), width=Inches(5.5))
                    # Caption
                    if alt:
                        cap = doc.add_paragraph(f"图：{alt}")
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap.runs[0].font.size = Pt2(10) if cap.runs else None
                except Exception:
                    p = doc.add_paragraph(f"[图片: {alt}]")
            else:
                p = doc.add_paragraph(f"[图片缺失: {alt}]")
            i += 1
            continue

        # 🆕 Table detection (simple pipe tables)
        if '|' in stripped and stripped.count('|') >= 2:
            # Check if next line is a separator line (|---|---|)
            if i + 2 < len(lines) and re.match(r'^[\s\|:\-]+$', lines[i + 1].strip()):
                # It's a table — collect rows
                table_rows = []
                header_cells = [c.strip() for c in stripped.split('|') if c.strip()]
                if header_cells:
                    table_rows.append(header_cells)
                j = i + 2  # skip separator
                while j < len(lines):
                    row_line = lines[j].strip()
                    if '|' in row_line and row_line.count('|') >= 2:
                        cells = [c.strip() for c in row_line.split('|') if c.strip()]
                        if cells:
                            table_rows.append(cells)
                        j += 1
                    else:
                        break

                if table_rows:
                    max_cols = max(len(r) for r in table_rows)
                    tbl = doc.add_table(rows=len(table_rows), cols=max_cols, style='Table Grid')
                    for ri, row in enumerate(table_rows):
                        for ci, cell_text in enumerate(row):
                            if ci < max_cols:
                                tbl.cell(ri, ci).text = cell_text
                    doc.add_paragraph("")  # spacer
                    i = j
                    continue

        # Headings
        heading_found = False
        for lv, prefix in [(1, '# '), (2, '## '), (3, '### ')]:
            if stripped.startswith(prefix):
                doc.add_heading(stripped[len(prefix):], level=lv)
                heading_found = True
                break
        if heading_found:
            i += 1
            continue

        # Regular paragraph — clean LaTeX delimiters
        text = stripped
        # Keep $...$ for readability (they would be OMML in pandoc mode)
        text = re.sub(r'\$\$([^$]+?)\$\$', r'[\1]', text)
        text = re.sub(r'\$([^$]+?)\$', r'\1', text)
        p = doc.add_paragraph(text)
        i += 1

    # Flush any remaining code block
    if in_code_block:
        _flush_code_block()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return str(output_path)
